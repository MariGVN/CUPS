from flask import Flask, render_template, request, send_file
import pandas as pd
import os
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys
from random import choice

app = Flask(__name__)
resultados_guardados = []
resultados_busqueda = []

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
def generar_analisis(resultados):
    patrones = {}
    for resultado in resultados:
        if not resultado["encontrado"]:
            continue
        for coincidencia in resultado["coincidencias"]:
            datos = coincidencia["datos"]
            for campo, valor in datos.items():
                if campo not in patrones:
                    patrones[campo] = {}
                if valor not in patrones[campo]:
                    patrones[campo][valor] = 0
                patrones[campo][valor] += 1
    resumen = []
    for campo, valores in patrones.items():
        for valor, cantidad in valores.items():
            if cantidad > 1:
                resumen.append({
                    "campo": campo,
                    "valor": valor,
                    "cantidad": cantidad
                })              
    return resumen
def detectar_patrones(resultados):
    patrones = {}
    for resultado in resultados:
        cups = resultado["cups"]
        if not resultado["encontrado"]:
            continue
        for coincidencia in resultado["coincidencias"]:
            datos = coincidencia["datos"]
            for campo, valor in datos.items():
                
                if campo not in patrones:
                    patrones[campo] = {}
                if valor not in patrones[campo]:
                    patrones[campo][valor] = {}
                
                if cups not in patrones[campo][valor]:
                    patrones[campo][valor][cups] = []
                patrones[campo][valor][cups].append({
                    "archivo": coincidencia["archivo"],
                    "hoja": coincidencia["hoja"],
                    "datos": coincidencia["datos"]
                })
    patrones_detectados = []
    for campo, valores in patrones.items():
        for valor, cups_dict in valores.items():
            if len(cups_dict) > 1:
                patrones_detectados.append({
                    "campo": campo,
                    "valor": valor,
                    "cantidad": len(cups_dict),
                    "cups": cups_dict
                })                    
    patrones_detectados.sort(
        key=lambda x: x["cantidad"],
        reverse=True
    )      
    estadisticas = {}
    if patrones_detectados:
        patron_principal = patrones_detectados[0]
        estadisticas = {
            "total_patrones": len(patrones_detectados),
            "campo_mas_repetido": patron_principal["campo"],
            "valor_mas_repetido": patron_principal["valor"],
            "cantidad_maxima": patron_principal["cantidad"]
        }
        conclusion = (
            f"Se detectaron {len(patrones_detectados)} patrones,"
            f"el patrón más frecuente corresponde al campo "
            f"{patron_principal['campo']} con el valor "
            f"{patron_principal['valor']}, "
            f"Compartido por {patron_principal['cantidad']} CUPS"
        )
    else:
        conclusion = (
            "No se detectaron patrones compartidos "
            "entre los CUPS analizados"
        )
    return{
        "patrones": patrones_detectados,
        "estadisticas": estadisticas,
        "conclusion": conclusion,
        "total_cups": len(resultados)
    }
   
@app.route("/")
def inicio():

    archivos = []

    for archivo in os.listdir(UPLOAD_FOLDER):

        if archivo.endswith(".xlsx"):

            hojas = []

            try:
                with pd.ExcelFile(
                    os.path.join(
                        UPLOAD_FOLDER,
                        archivo
                    )
                ) as excel:
                    hojas = excel.sheet_names
            except Exception:
                pass

                hojas = excel.sheet_names

            except Exception:
                pass

            archivos.append({
                "nombre": archivo,
                "hojas": hojas
            })

    return render_template(
        "buscador.html",
        archivos=archivos
    )

@app.route("/subir", methods=["POST"])
def subir():

    archivo = request.files["excel"]

    if archivo.filename:

        ruta = os.path.join(
            UPLOAD_FOLDER,
            archivo.filename
        )

        archivo.save(ruta)

    return inicio()

@app.route("/buscar", methods=["POST"])
def buscar():

    cups_buscado = (
        request.form["cups"]
        .strip()
        .upper()
    )

    resultados = []

    for archivo in os.listdir(UPLOAD_FOLDER):

        if not archivo.endswith(".xlsx"):
            continue

        ruta = os.path.join(
            UPLOAD_FOLDER,
            archivo
        )

        try:

            hojas = pd.read_excel(
                ruta,
                sheet_name=None
            )

            for nombre_hoja, df in hojas.items():

                try:

                    # Caso especial datos generales
                    if nombre_hoja.upper() == "DATOS GENERALES":

                        df = pd.read_excel(
                            ruta,
                            sheet_name=nombre_hoja,
                            header=2
                        )
                    # Caso especial VF ELECTROMECANICOS ELIMINAR
                    if nombre_hoja.upper() == "VF ELECTROMECANICOS ELIMINAR":

                        df = pd.read_excel(
                            ruta,
                            sheet_name=nombre_hoja,
                            header=2
                        )

                    columna_cups = next(
                        (
                            col
                            for col in df.columns
                            if "CUPS" in str(col).upper()
                        ),
                        None
                    )

                    if columna_cups is None:
                        continue

                    df[columna_cups] = (
                        df[columna_cups]
                        .astype(str)
                        .str.strip()
                        .str.upper()
                    )

                    coincidencias = (
                        df[df[columna_cups] == cups_buscado]
                        
                    )

                    for _, fila in coincidencias.iterrows():

                        datos_fila = {}

                        for columna, valor in fila.items():

                            if pd.isna(valor):
                                continue

                            valor = str(valor).strip()

                            if valor == "":
                                continue

                            datos_fila[columna] = valor

                        resultados.append({
                            "archivo": archivo,
                            "hoja": nombre_hoja,
                            "datos": datos_fila
                        })

                except Exception as e:
                    print(
                        f"Error hoja {nombre_hoja}: {e}"
                    )

        except Exception as e:
            print(
                f"Error archivo {archivo}: {e}"
            )
    global resultados_busqueda
    resultados_busqueda = resultados
    return render_template(
        "resultados.html",
        cups=cups_buscado,
        resultados=resultados
    )
@app.route("/eliminar/<nombre>")
def eliminar(nombre):

    ruta = os.path.join(
        UPLOAD_FOLDER,
        nombre
    )

    if os.path.exists(ruta):
        os.remove(ruta)

    return inicio()
@app.route("/eliminar_todos")
def eliminar_todos():

    for archivo in os.listdir(UPLOAD_FOLDER):

        if archivo.endswith(".xlsx"):

            ruta = os.path.join(
                UPLOAD_FOLDER,
                archivo
            )

            if os.path.exists(ruta):
                os.remove(ruta)

    return inicio()
@app.route("/analisis")
def pagina_analisis():
    return render_template(
        "analisis.html"
    )
@app.route("/deteccion_de_patrones")
def deteccion_de_patrones():
    analisis_patrones = detectar_patrones(
        resultados_guardados
    )
    return render_template(
        "deteccion_de_patrones.html",
        analisis_patrones=analisis_patrones
    )
@app.route("/analizar", methods=["POST"])
def analizar():

    cups_texto = request.form["cups"]

    lista_cups = [
        c.strip().upper()
        for c in cups_texto.splitlines()
        if c.strip()
    ]

    # Creacion de resultados una sola vez
    resultados_dict = {}

    for cups in lista_cups:
        resultados_dict[cups] = {
            "cups": cups,
            "encontrado": False,
            "coincidencias": []
        }

    # Recorre excel una vez
    for archivo in os.listdir(UPLOAD_FOLDER):

        if not archivo.endswith(".xlsx"):
            continue

        ruta = os.path.join(
            UPLOAD_FOLDER,
            archivo
        )

        try:

            hojas = pd.read_excel(
                ruta,
                sheet_name=None
            )

            for nombre_hoja, df in hojas.items():

                try:

                    # Hojas especiales (no empiezan como las demas hojas)
                    if nombre_hoja.upper() == "DATOS GENERALES":

                        df = pd.read_excel(
                            ruta,
                            sheet_name=nombre_hoja,
                            header=2
                        )

                    if nombre_hoja.upper() == "VF ELECTROMECANICOS ELIMINAR":

                        df = pd.read_excel(
                            ruta,
                            sheet_name=nombre_hoja,
                            header=2
                        )

                    columna_cups = next(
                        (
                            col
                            for col in df.columns
                            if "CUPS" in str(col).upper()
                        ),
                        None
                    )

                    if columna_cups is None:
                        continue

                    df[columna_cups] = (
                        df[columna_cups]
                        .astype(str)
                        .str.strip()
                        .str.upper()
                    )

                    # Buscar todos los CUPS en la hoja
                    for cups_buscado in lista_cups:

                        filas = df[
                            df[columna_cups] == cups_buscado
                        ]

                        if filas.empty:
                            continue

                        resultados_dict[cups_buscado]["encontrado"] = True

                        for _, fila in filas.iterrows():

                            datos = {}

                            for columna, valor in fila.items():

                                if pd.isna(valor):
                                    continue

                                valor = str(valor).strip()

                                if valor == "":
                                    continue

                                datos[columna] = valor

                            resultados_dict[cups_buscado]["coincidencias"].append({
                                "archivo": archivo,
                                "hoja": nombre_hoja,
                                "datos": datos
                            })

                except Exception as e:

                    print(
                        f"Error hoja {nombre_hoja}: {e}"
                    )

        except Exception as e:

            print(
                f"Error archivo {archivo}: {e}"
            )

    resultados_analisis = list(
        resultados_dict.values()
    )

    global resultados_guardados
    resultados_analisis = list(
        resultados_dict.values()
    )
    resultados_guardados = resultados_analisis
    resumen = generar_analisis(
        resultados_analisis
    )
    return render_template(
        "analisis_resultados.html",
        resultados=resultados_analisis,
        analisis=resumen
    )
@app.route("/exportar_excel")
def exportar_excel():
    resumen = []
    detalles = []
    

    for resultado in resultados_guardados:
        cups = resultado["cups"]
        cantidad = len(
            resultado["coincidencias"]
        )
        resumen.append({
            "CUPS": cups,
            "Apariciones": cantidad
        })
        for coincidencia in resultado["coincidencias"]:
            detalles.append({
                "CUPS": cups,
                "ID ACTIVIDAD (Nº TDC)": coincidencia["datos"].get("ID ACTIVIDAD (Nº TDC)", ""),
                "Archivo": coincidencia["archivo"],
                "Hoja": coincidencia["hoja"]
            })
            
    ruta_excel = os.path.join(
        os.getcwd(),
        "analisis_CUPS.xlsx"
    )

    with pd.ExcelWriter(ruta_excel, engine="openpyxl") as writer:
        ahora = datetime.now()
        info = pd.DataFrame([
            ["Informe de análisis de CUPS", ""],
            ["Fecha", ahora.strftime("%d/%m/%Y")],
            ["Hora", ahora.strftime("%H:%M:%S")],
            ["Total de CUPS analizados", len(resultados_guardados)],
            ["Total de coincidencias", sum(len(r["coincidencias"]) for r in resultados_guardados)],
            ["Versión", "1.0"],
            ["", ""],
            ["Generado automáticamente por la aplicación de generador de CUPS.", ""]
        ])
        info.to_excel(
            writer,
            sheet_name="Información", 
            index=False,
            header=False
        )
        pd.DataFrame(resumen).to_excel(
            writer,
            sheet_name="Resumen Ejecutivo",
            index=False
        )
        pd.DataFrame(detalles).to_excel(
            writer,
            sheet_name="Coincidencias",
            index=False
        )
        hoja_datos = writer.book.create_sheet("Datos Encontrados")
        fila = 1
        borde = Border(
                    left=Side(style="thin", color="000000"),
                    right=Side(style="thin", color="000000"),
                    top=Side(style="thin", color="000000"),
                    bottom=Side(style="thin", color="000000")
                )
        azul = PatternFill(
            fill_type="solid",
            start_color="1F4E79"
        )
        gris = PatternFill(
            fill_type="solid",
            start_color="DCE6F1"
        )
        colores_disponibles = [
            "1F4E79", "C00000", "70AD47", "ED7D31", "7030A0",
            "00B0F0", "BF9000", "FF66CC", "548235", "002060",
            "A52A2A", "5B9BD5", "FFC000", "92D050", "8064A2",
            "E26B0A", "31869B", "953735", "76923C", "4F81BD"
        ]
        color_por_cups = {}
        for i, resultado in enumerate(resultados_guardados):
            cups = resultado["cups"]
            color_por_cups[cups] = colores_disponibles[
                i % len(colores_disponibles)
            ]
            relleno_cups = PatternFill(
                fill_type="solid",
                start_color=color_por_cups[cups]
            )
            cups = resultado["cups"]
            for coincidencia in resultado["coincidencias"]:
                hoja_datos.merge_cells(
                    start_row=fila,
                    start_column=1,
                    end_row=fila,
                    end_column=2
                )
                hoja_datos.cell(fila,1).value = f"CUPS: {cups}"
                hoja_datos.cell(fila,1).font = Font(
                    bold=True,
                    size=14,
                    color="FFFFFF"
                )
                hoja_datos.cell(fila,1).fill = relleno_cups
                hoja_datos.cell(fila,1).alignment = Alignment(
                    horizontal="center"
                )
                fila += 2
                hoja_datos.cell(fila,1).value = "Archivo"
                hoja_datos.cell(fila,1).font = Font(bold=True)
                hoja_datos.cell(fila,1).fill = gris
                hoja_datos.cell(fila,1).border = borde

                hoja_datos.cell(fila,2).value = coincidencia["archivo"]
                hoja_datos.cell(fila,2).border = borde

                fila += 1
                hoja_datos.cell(fila,1).value = "Hoja"
                hoja_datos.cell(fila,1).font = Font(bold=True)
                hoja_datos.cell(fila,1).fill = gris
                hoja_datos.cell(fila,1).border = borde

                hoja_datos.cell(fila,2).value = coincidencia["hoja"]
                hoja_datos.cell(fila,2).border = borde
                fila += 2
                hoja_datos.cell(fila,1).value = "Campo"
                hoja_datos.cell(fila,2).value = "Valor"
                for columna in  [1,2]:
                    hoja_datos.cell(fila,columna).font = Font(
                        bold=True,
                        color="FFFFFF"
                    )
                    hoja_datos.cell(fila,columna).fill = relleno_cups
                    hoja_datos.cell(fila,columna).alignment = Alignment(
                        horizontal="center"
                    )
                    hoja_datos.cell(fila,columna).border = borde
                fila += 1
                for campo, valor in coincidencia["datos"].items():
                    hoja_datos.cell(fila,1).value = campo
                    hoja_datos.cell(fila,2).value = valor
                    hoja_datos.cell(fila,1).border = borde
                    hoja_datos.cell(fila,2).border = borde
                    hoja_datos.cell(fila,1).alignment = Alignment(
                        wrap_text=True,
                        vertical="top"
                    )
                    hoja_datos.cell(fila,2).alignment = Alignment(
                        wrap_text=True,
                        vertical="top"
                    )
                    texto = str(valor)
                    lineas = max(1, len(texto) // 45 + 1)
                    hoja_datos.row_dimensions[fila].height = lineas * 18
                    fila += 1
                fila += 2
        hoja_datos.column_dimensions["A"].width = 40
        hoja_datos.column_dimensions["B"].width = 35                          
        for fila in hoja_datos.iter_rows():
            for celda in fila:
                if celda.value is not None:
                    celda.alignment = Alignment(
                        horizontal=celda.alignment.horizontal,
                        vertical="top",
                        wrap_text= True
                    )
        for hoja in writer.book.worksheets:
            if hoja.title == "Información":
                hoja.merge_cells("A1:B1")
                hoja["A1"].font = Font(
                    bold=True,
                    size=18,
                    color="FFFFFF"
                )
                hoja["A1"].fill = PatternFill(
                    fill_type="solid",
                    start_color="1F4E79"
                )
                hoja["A1"].alignment = Alignment(
                    horizontal="center",
                    vertical="center"
                )
                hoja.row_dimensions[1].height = 30
                relleno = PatternFill(
                        fill_type="solid",
                        start_color="DCE6F1"
                    )
                
                for fila in [2, 3, 4, 5, 6]:
                    hoja[f"A{fila}"].font = Font(
                        name="Arial",
                        size=11,
                        bold=True
                    )
                    hoja[f"A{fila}"].fill = relleno
                    hoja[f"A{fila}"].border = borde
                    hoja[f"B{fila}"].border = borde
                for fila in [4, 5]:
                    hoja[f"B{fila}"].alignment = Alignment(horizontal="center")
                
                hoja["A8"].font = Font(
                    italic=True,
                    color="808080"
                ) 
                hoja.column_dimensions["A"].width = 40
                hoja.column_dimensions["B"].width = 30
                for fila in hoja.iter_rows():
                    for celda in fila:
                        if celda.value is not None:
                            celda.border = borde
                            celda.alignment = Alignment(
                                horizontal=celda.alignment.horizontal,
                                vertical="top",
                                wrap_text=True
                            )

                continue
            if hoja.title != "Datos Encontrados":
                for columna in hoja.columns:
                    longitud = max(
                        len(str(c.value)) if c.value is not None else 0
                        for c in columna
                    )
                    hoja.column_dimensions[
                        get_column_letter(columna[0].column)
                    ].width = longitud + 3
            for celda in hoja[1]:
                celda.font = Font(bold=True, color="FFFFFF")
                celda.fill = PatternFill(
                    fill_type="solid",
                    start_color="1F4E79"
                )
                celda.alignment = Alignment(horizontal="center")
            
            hoja.auto_filter.ref = hoja.dimensions
            hoja.freeze_panes = "A2"
            for fila in hoja.iter_rows():
                for celda in fila:
                    if celda.value is not None:
                        celda.border = borde
    return send_file(
        ruta_excel,
        as_attachment=True
    )
@app.route("/exportar_word")
def exportar_word():
   analisis_patrones = detectar_patrones(resultados_guardados)
   documento = Document()
   titulo = documento.add_heading(
       "Informe de detección de patrones",
       level=1
   )
   titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
   subtitulo = documento.add_paragraph(
       "Sistema de análisis automático de CUPS"
   )
   subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
   documento.add_paragraph()
   fecha = documento.add_paragraph(
       f"Fecha: {datetime.now().strftime('%d/%m/%Y')}"
   )
   fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
   hora = documento.add_paragraph(
       f"Hora: {datetime.now().strftime('%H:%M:%S')}"
   )
   hora.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
   documento.add_page_break()
   documento.add_heading(
       "1. Resumen Ejecutivo",
       level=1
   )
   documento.add_paragraph(
       "Se realizó un análisis automático de los CUPS "
       "proporcionados con el objetivo de detectar "
       "patrones compartidos entre los diferentes "
       "archivos Excel cargados en la aplicación. "
   )
   estadisticas = analisis_patrones["estadisticas"]
   documento.add_paragraph(
       f"Total de CUPS analizados: "
       f"{analisis_patrones['total_cups']}"
   )
   documento.add_paragraph(
       f"Patrones detectados: "
       f"{estadisticas.get('total_patrones',0)}"
   )
   documento.add_heading(
       "2. Estadísticas Generales",
       level=1
   )
   tabla = documento.add_table(rows=1, cols=2)

   tabla.style = "Table Grid"

   encabezado = tabla.rows[0].cells

   encabezado[0].text = "Indicador"

   encabezado[1].text = "Valor"
   datos = [
       (
           "Total de cups analizados",
           str(
               analisis_patrones["total_cups"]
           )
       ),
       (
           "Patrones detectados",
           str(
               estadisticas.get(
                   "total_patrones",
                   0
               )
           )
       ),
       (
           "Campo más repetido",
           estadisticas.get(
               "valor más repetido",
               "-"
           )
       ),
       (
           "Valor más repetido",
           estadisticas.get(
               "valor_mas_repetido",
               "-"
           )
       ),
       (
           "Mayor número de coincidencias",
           str(
               estadisticas.get(
                   "cantidad_maxima",
                   0
               )
           )
       )
   ]
   for indicador, valor in datos:
       fila = tabla.add_row().cells
       fila[0].text = indicador
       fila[1].text = valor
   documento.add_heading(
       "3. Patrones detectados",
       level=1
   )
   if not analisis_patrones["patrones"]:
       documento.add_paragraph(
           "No se detectaron patrones compartidos entre los CUPS analizados."
       )
   else:
       for i,patron in enumerate(analisis_patrones["patrones"], start=1):
           documento.add_heading(
               f"Patron {i}",
               level=2
           )
           documento.add_paragraph(
               f"Campo: {patron['campo']}"
           )
           documento.add_paragraph(
               f"Valor: {patron['valor']}"
           )
           documento.add_paragraph(
               f"Compartido por: {patron['cantidad']} CUPS"
           )
           documento.add_paragraph(
               "CUPS implicados: "
           )
           for cups in patron["cups"]:
               documento.add_paragraph(
                   cups,
                   style="List Bullet"
               )
           documento.add_paragraph(
               "Ubicación de las coincidencias:"
           )
           tabla_patron = documento.add_table(
               rows=1,
               cols=4
           )
           tabla_patron.style = "Table Grid"
           encabezado = tabla_patron.rows[0].cells
           encabezado[0].text = "CUPS"
           encabezado[1].text = "ID ACTIVIDAD (Nº TDC)"
           encabezado[2].text = "Archivo"
           encabezado[3].text = "Hoja"
           for cups, coincidencias in patron ["cups"].items():
               for coincidencia in coincidencias:
                   fila = tabla_patron.add_row().cells
                   fila[0].text = cups
                   fila[1].text = coincidencia.get("datos", {}).get("ID ACTIVIDAD (Nº TDC)", "")
                   fila[2].text = coincidencia["archivo"]
                   fila[3].text = coincidencia["hoja"]
           documento.add_paragraph
   documento.add_heading(
        "4. Conclusiones",
        level=1
    )
   documento.add_paragraph(
       analisis_patrones["conclusion"]
   )
   documento.add_paragraph()
   documento.add_paragraph(
       "Documento generado automáticamente por la aplicación."
   )
   ruta = os.path.join(
       os.getcwd(),
       "informe_patrones.docx"
   )
   documento.save(ruta)
   
   return send_file(
        ruta,
        as_attachment=True
    )
@app.route("/exportar_excel_busqueda")
def exportar_excel_busqueda():
    ruta_excel = os.path.join(
        os.getcwd(),
        "exportar_excel_busqueda_CUPS.xlsx"
    )
    with pd.ExcelWriter(ruta_excel, engine="openpyxl") as writer:
        pd.DataFrame().to_excel(
            writer,
            sheet_name="Resultados",
            index=False
        )
        hoja = writer.book["Resultados"]
        hoja.delete_rows(1)
        
        borde = Border(
            left=Side(style="thin", color="000000"),
            right=Side(style="thin", color="000000"),
            top=Side(style="thin", color="000000"),
            bottom=Side(style="thin", color="000000"),
        )
        azul = PatternFill(
            fill_type="solid",
            start_color="1F4E79"
        )
        gris = PatternFill(
            fill_type="solid",
            start_color="DCE6F1"
        )
        fila = 1
        for coincidencia in resultados_busqueda:
            hoja.merge_cells(f"A{fila}:B{fila}")
            hoja[f"A{fila}"] = (
                f"CUPS: {coincidencia['datos'].get('CUPS', '')}"
            )
            hoja[f"A{fila}"].font =Font(
                bold=True,
                size=14,
                color="FFFFFF"
            )
                       
            hoja[f"A{fila}"].fill = azul
            hoja[f"A{fila}"].alignment = Alignment(
                horizontal="center"
            )
            fila += 2
            
            hoja[f"A{fila}"] = "Archivo"
            hoja[f"A{fila}"].font = Font(bold=True)
            hoja[f"A{fila}"].fill = gris
            hoja[f"A{fila}"].border = borde

            hoja[f"B{fila}"] = coincidencia["archivo"]
            hoja[f"B{fila}"].border = borde
            
            fila += 1
            hoja[f"A{fila}"] = "Hoja"
            hoja[f"A{fila}"].font = Font(bold=True)
            hoja[f"A{fila}"].fill = gris
            hoja[f"A{fila}"].border = borde

            hoja[f"B{fila}"] = coincidencia["hoja"]
            hoja[f"B{fila}"].border = borde
        
            fila += 2
            hoja[f"A{fila}"] = "Campo"
            hoja[f"B{fila}"] = "Valor"
           
            for celda in [f"A{fila}", f"B{fila}"]:
                hoja[celda].font = Font(
                    bold=True,
                    color="FFFFFF"
                )
                hoja[celda].fill = azul
                hoja[celda].alignment = Alignment(horizontal="center")
                hoja[celda].border = borde
            
            fila += 1
            for campo, valor in coincidencia["datos"].items():
                hoja[f"A{fila}"] = campo
                hoja[f"B{fila}"] = valor
                hoja[f"A{fila}"].border = borde
                hoja[f"B{fila}"].border = borde
                fila += 1
            fila += 2
        hoja.column_dimensions["A"].width = 40
        hoja.column_dimensions["B"].width = 60
        for fila in hoja.iter_rows():
            for celda in fila:
                if celda.value is not None:
                    celda.alignment = Alignment(
                        horizontal=celda.alignment.horizontal,
                        vertical="top",
                        wrap_text=True
                    )
    return send_file(
        ruta_excel,
        as_attachment=True
    )

           


if __name__ == "__main__":
    app.run(debug=True)


